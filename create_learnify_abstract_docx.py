import os
from docx import Document
from docx.shared import Pt


def build_document(doc: Document) -> None:
    title = "Learnify — A Personalized Microlearning Platform for Time-Bound Learning Goals"
    doc.add_heading(title, 0)

    doc.add_heading("Abstract — Personalized Microlearning and Scheduling Web Application", level=1)

    paragraphs = [
        (
            "Learnify is a web-based microlearning application designed to make long-term learning structured, "
            "interactive, and efficient. The system allows users to enter their learning goal (e.g., UPSC, GRE, "
            "Programming, etc.) and target duration (for example, one year). Based on this input, Learnify automatically "
            "breaks the subject into micro-topics and generates a personalized study plan with daily or weekly learning "
            "goals, reminders, and quizzes to enhance retention and motivation."
        ),
        (
            "The main problem addressed is that learners often struggle to stay consistent and manage vast syllabi due "
            "to poor time management and lack of structure. Learnify solves this by dividing subjects into smaller, "
            "achievable milestones with periodic interactive quizzes, notifications, and progress tracking. The platform "
            "uses adaptive algorithms to adjust topic difficulty and reminder frequency according to user performance "
            "and schedule adherence."
        ),
        (
            "Technically, Learnify follows a client–server architecture. The frontend is built using HTML, CSS, and "
            "JavaScript, offering a responsive and distraction-free UI optimized for both desktop and mobile users. "
            "The backend, powered by Python (Flask or Django), handles user registration, schedule generation, quiz "
            "management, and performance analytics. The database (SQLite or PostgreSQL) stores users, subjects, topics, "
            "quiz questions, progress logs, and reminder schedules."
        ),
        (
            "Learnify also integrates email or push notifications for reminders, gamified achievements (like badges for "
            "consistency), and AI-assisted topic breakdown to simplify large subjects into digestible learning units. "
            "Security measures include user authentication, secure session management, and sanitized input handling."
        ),
    ]

    for text in paragraphs:
        para = doc.add_paragraph(text)
        para.paragraph_format.space_after = Pt(12)

    label = doc.add_paragraph("Expected outcomes:")
    for run in label.runs:
        run.bold = True

    outcomes = [
        "Improved time management and retention through microlearning.",
        "Increased learner engagement via interactive quizzes and reminders.",
        "Scalable web architecture suitable for schools, universities, and self-learners.",
    ]

    for bullet in outcomes:
        doc.add_paragraph(bullet, style="List Bullet")

    closing = (
        "The project demonstrates the power of structured microlearning, combining intelligent scheduling, "
        "gamification, and personalized analytics to make continuous learning effective and enjoyable."
    )
    doc.add_paragraph(closing)


def main() -> None:
    output_dir = "/workspace/docs"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Learnify_Abstract.docx")

    document = Document()
    build_document(document)
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
